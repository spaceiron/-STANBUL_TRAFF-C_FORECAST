#!/usr/bin/env python3
"""Patch TRAFİKPULS.docx on Desktop → TRAFİKPULS_FIXED.docx"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DESKTOP = Path('/Users/uzaydemir/Desktop')

SECTION_9_BLURBS = {
    '9.1) Login and Registration Screens': (
        'This screen is the entry point of the application. Users sign in with email and '
        'password through Firebase Authentication. Successful login opens the map interface.'
    ),
    '9.2) Map and Route Search Screen': (
        'Users search lines such as M4 and view stop markers with route polylines. '
        'Stop coordinates and ETA values are loaded from the Flask API.'
    ),
    '9.3) Prediction Bottom Sheet (Density, Confidence, Forecast)': (
        'The draggable panel shows density score, confidence level, 30/60-minute forecast, '
        'incidents, and alternative routes when crowding is high.'
    ),
    '9.4) Feedback Screen': (
        'Passengers report empty, standing, or full occupancy. Data is saved to Firestore '
        'and sent to /report_density to update live predictions.'
    ),
    '9.5) Profile and Settings Screen': (
        'Profile settings include language (TR/EN), favorite routes, and a density alert '
        'threshold for local notifications.'
    ),
    '9.6) Admin Dashboard Screen': (
        'Admin-lite view shows feedback stats, model metrics, and prototype retrain/push actions.'
    ),
}

LIST_OF_FIGURES_LINES = [
    'LIST OF FIGURES',
    '',
    'Figure 5.1 – High-level architecture of the proposed TrafikPuls system',
    'Figure 7.1 – Context diagram of TrafikPuls',
    'Figure 7.2 – Use case diagram of TrafikPuls',
    'Figure 7.3 – Activity diagram (route selection flow)',
    'Figure 7.4 – Sequence diagram (prediction and feedback flow)',
    'Figure 7.5 – Simplified class diagram (mobile layer)',
    'Figure 7.6 – Component diagram of TrafikPuls',
    'Figure 7.7 – Level 0 data flow diagram (DFD)',
    'Figure 9.1 – Login and registration screen',
    'Figure 9.2 – Map screen with route search',
    'Figure 9.3 – Prediction bottom sheet',
    'Figure 9.4 – Crowdsourcing feedback screen',
    'Figure 9.5 – Profile and settings screen',
    'Figure 9.6 – Admin dashboard screen',
]


def find_docx() -> Path:
    for p in DESKTOP.iterdir():
        if p.suffix.lower() != '.docx' or 'PULS' not in p.name.upper():
            continue
        if 'FIXED' in p.name.upper() or 'ACCEPTANCE' in p.name.upper():
            continue
        return p
    raise FileNotFoundError('TRAFİKPULS.docx not found on Desktop')


def delete_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def is_duplicate_toc_block(texts: list[str], start: int) -> bool:
    if start >= len(texts) or texts[start] != '1) Introduction':
        return False
    for j in range(start, min(start + 50, len(texts))):
        if texts[j] == '11) References':
            return True
    return False


def split_merged_uml(paragraph) -> None:
    text = paragraph.text.strip()
    if not text.startswith('7.') or 'The ' not in text or len(text) < 70:
        return
    m = re.match(r'^(7\.\d[^a-zA-Z]*)\s*([^\n]+?)\s*(The .+)$', text, re.S)
    if not m:
        return
    heading = m.group(1).strip()
    if not heading.endswith(')'):
        heading = heading.rstrip() + ')'
    body = m.group(3).strip()
    paragraph.text = heading + ' ' + m.group(2).strip().split('\n')[0]
    insert_after(paragraph, body)


def main() -> None:
    src = find_docx()
    out = DESKTOP / 'TRAFİKPULS_FIXED.docx'
    shutil.copy2(src, out)
    doc = Document(str(out))

    texts = [p.text.strip() for p in doc.paragraphs]
    to_delete = []
    for i in range(len(texts)):
        if is_duplicate_toc_block(texts, i):
            j = i
            while j < len(texts) and texts[j] != 'INTRODUCTION':
                to_delete.append(j)
                j += 1
    for i in sorted(set(to_delete), reverse=True):
        delete_paragraph(doc.paragraphs[i])

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == 'Uzay Demir':
            nxt = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ''
            if nxt != '210218027':
                insert_after(p, '210218027')
            break

    for p in doc.paragraphs:
        if p.text.startswith('Our esteemed professor'):
            p.text = (
                'I would like to express my sincere gratitude to my project supervisor, '
                'Dr. JOHN OLORUNFEMI OLAIFA, for his guidance and support throughout '
                'this graduation project.'
            )

    seen_22 = False
    seen_51 = False
    seen_9: set[str] = set()
    paras = list(doc.paragraphs)
    for idx, p in enumerate(paras):
        t = p.text.strip()
        if t == '2.2) Existing System with Limitations' and not seen_22:
            insert_after(p, 'A summary comparison is presented in Table 2.1.')
            seen_22 = True
        if t == '5.1) Advantages of the Proposed System' and not seen_51:
            insert_after(p, 'The overall system interaction is illustrated in Figure 5.1.')
            seen_51 = True
        if t in SECTION_9_BLURBS and t not in seen_9:
            nxt = paras[idx + 1].text.strip() if idx + 1 < len(paras) else 'x'
            if len(nxt) < 25:
                insert_after(p, SECTION_9_BLURBS[t])
                seen_9.add(t)
        split_merged_uml(p)

    for p in doc.paragraphs:
        if p.text.strip().startswith('Keywords:'):
            anchor = p
            for line in reversed(LIST_OF_FIGURES_LINES):
                anchor = insert_after(anchor, line)
            break

    doc.save(str(out))
    print(f'OK → {out}')
    print('Word: References → Update Table of Contents → captions on images → Export PDF')


if __name__ == '__main__':
    main()
